from flask import Flask, render_template, request, send_file, jsonify
import json
import os
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

app = Flask(__name__)


# Load offices from JSON file
def load_offices():
    try:
        with open("officers.json", "r", encoding="utf-8") as file:
            data = json.load(file)
            return data.get("offices", [])
    except Exception as e:
        return ["कोई कार्यालय उपलब्ध नहीं"]


@app.route("/")
def index():
    offices = load_offices()
    return render_template("index.html", offices=offices)


@app.route("/save_transcript", methods=["POST"])
def save_transcript():
    try:
        data = request.json
        heading = data.get("heading", "").strip()
        transcript = data.get("transcript", "").strip()
        office = data.get("office", "").strip()

        # Create a Word document
        doc = Document()

        # Set font to Noto Sans Devanagari for Hindi text
        def set_font(paragraph, text):
            run = paragraph.add_run(text)
            run.font.name = "Noto Sans Devanagari"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Devanagari")
            run.font.size = Pt(12)

        # Add content to the document
        set_font(doc.add_paragraph(), f"शीर्षक: {heading}")
        doc.add_paragraph()  # Empty line
        set_font(doc.add_paragraph(), f"संबंधित कार्यालय: {office}")
        doc.add_paragraph()  # Empty line
        set_font(doc.add_paragraph(), "प्रतिलेख:")
        set_font(doc.add_paragraph(), transcript)

        # Save document to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name="transcript.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return jsonify({"error": f"फ़ाइल सहेजने में त्रुटि: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(debug=True)
